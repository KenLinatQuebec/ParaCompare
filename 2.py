import os
import difflib
from docx import Document
from docx.shared import RGBColor
from PIL import Image
import pytesseract
from split_run import split_run_by, insert_run_before, insert_run_after, copy_run_format


def ocr_function(image_path):
    # 图像 OCR 转换
    try:
        image = Image.open(image_path)
        ocr_text = pytesseract.image_to_string(image, lang="chi_sim", config="--psm 4")
    except FileNotFoundError:
        print("图像文件未找到，请检查路径是否正确。")
        exit()
    except Exception as e:
        print("图像 OCR 转换出错:", e)
        exit()
    return ocr_text


def ocr_table(image_path):
    try:
        image = Image.open(image_path)
        ocr_text = pytesseract.image_to_string(image, lang="eng", config="--psm 6")
    except FileNotFoundError:
        print("图像文件未找到，请检查路径是否正确。")
        exit()
    except Exception as e:
        print("图像 OCR 转换出错:", e)
        exit()
    return ocr_text


def read_word_document(file_path):
    doc = Document(file_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + " "
    return text.strip()


def word_preprocess(text):
    return text.replace(" ", "").replace(",", "，").replace("\n", "").replace(":", "：").replace("“", "“").replace("!",
                                                                                                                 "！").replace(
        "(", "（")


def find_edit_positions(ori, ocr):
    matcher = difflib.SequenceMatcher(None, ori, ocr)
    edit_positions = []
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == 'insert':
            edit_positions.append(('insert', i1, ocr[j1:j2]))
        elif tag == 'delete':
            edit_positions.append(('delete', i1, ori[i1:i2]))
        elif tag == 'replace':
            # Include both the original text (that is being replaced)
            # and the new text in the replace action
            edit_positions.append(('replace', i1, ori[i1:i2], ocr[j1:j2]))
        elif tag == 'equal':
            edit_positions.append(('equal', i1, ori[i1:i2]))
    for action in edit_positions:
        if action[0] == 'replace':
            print(f"Replace '{action[2]}' with '{action[3]}' at position {action[1]} of the original text")
        else:
            print(f"{action[0].capitalize()} '{action[2]}' at position {action[1]} of the original text")
    return edit_positions


def mark_diff_preserve_format(doc, edit_positions):
    # Sort the edits in reverse order to avoid disrupting positions of subsequent edits
    edit_positions.sort(key=lambda x: x[1], reverse=True)

    for edit in edit_positions:
        action = edit[0]
        pos = edit[1]

        if action == 'replace':
            original_text = edit[2]
            new_text = edit[3]
        else:
            text = edit[2]

        for i, par in enumerate(doc.paragraphs):
            if len(par.text) <= pos:
                # Move on to the next paragraph if the position is not in the current one
                pos -= len(par.text)
                continue
            # Split the run containing the position
            for run in par.runs:
                if len(run.text) <= pos:
                    pos -= len(run.text)
                    continue
                runs = split_run_by(par, run, [pos])
                if action == 'insert':
                    # Insert text at the position
                    new_run = insert_run_before(par, runs[1], text)
                    # Copy formatting from original run to new run
                    copy_run_format(run, new_run)
                    # Mark the inserted text with red color
                    new_run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)  # RGB for red
                elif action == 'delete':
                    # Delete the text at the position
                    runs[1].text = text
                    # Mark the deleted text with strikethrough
                    runs[1].font.strike = True
                    runs[1].font.color.rgb = RGBColor(0xFF, 0x00, 0x00)  # RGB for red
                elif action == 'replace':
                    # Replace the original text with the new text
                    runs[1].text = original_text
                    # Mark the original text with strikethrough
                    runs[1].font.strike = True
                    runs[1].font.color.rgb = RGBColor(0xFF, 0x00, 0x00)  # RGB for red
                    # Insert the new text after the original text
                    new_run = insert_run_after(par, runs[1], new_text)
                    # Copy formatting from original run to new run
                    copy_run_format(run, new_run)
                    # Mark the new text with red color
                    new_run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)  # RGB for red
                break
            break
    return doc


def save_document(doc, modified_word_document_path):
    # 删除已经存在的同名文件
    if os.path.exists(modified_word_document_path):
        os.remove(modified_word_document_path)

    # 保存文档
    doc.save(modified_word_document_path)
    print("已保存修改后的 Word 文档:", modified_word_document_path)


def main(image_paths, word_document_paths, modified_word_document_path):
    for i in range(len(word_document_paths)):
        image_path = image_paths[i]
        word_document_path = word_document_paths[i]
        word_text = read_word_document(word_document_path)
        ocr_text = word_preprocess(ocr_function(image_path))
        table_ocr_text = ocr_table(image_path)  # 表格OCR识别文本
        positions = find_edit_positions(word_text, ocr_text)
        # 进行表格文本与原始文本的比较和标记处理
        table_positions = find_edit_positions(word_text, table_ocr_text)
        positions.extend(table_positions)  # 将表格的编辑位置合并到总的编辑位置列表中
        doc = Document(word_document_path)
        mark_diff_preserve_format(doc, positions)
        save_document(doc, modified_word_document_path)



if __name__ == '__main__':
    main(image_paths=["2.png"], word_document_paths=["2.docx"],modified_word_document_path="modified_document.docx")