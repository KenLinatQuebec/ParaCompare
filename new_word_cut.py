import os
from copy import deepcopy
from docx import Document


def insert_run_at_position(par, pos, txt=''):
    """Insert a new run with text {txt} into paragraph {par}
    at given position {pos}.
    Returns the newly created run.
    """
    p = par._p
    new_run = par.add_run(txt)
    p.insert(pos + 1, new_run._r)

    return new_run


def insert_run_before(par, run, txt=''):
    """Insert a new run with text {txt} into paragraph before given {run}.
    Returns the newly created run.
    """
    run_2 = par.add_run(txt)
    run._r.addprevious(run_2._r)

    return run_2


def insert_run_after(par, run, txt=''):
    """Insert a new run with text {txt} into paragraph after given {run}.
    Returns the newly created run.
    """
    run_2 = par.add_run(txt)
    run._r.addnext(run_2._r)

    return run_2


def copy_run_format(run_src, run_dst):
    """Copy formatting from {run_src} to {run_dst}.
    """
    rPr_target = run_dst._r.get_or_add_rPr()
    rPr_target.addnext(deepcopy(run_src._r.get_or_add_rPr()))
    run_dst._r.remove(rPr_target)


def split_run_by(par, run, split_by):
    """Split text in {run} from paragraph {par} by positions
    provided by {split_by}, while retaining original {run}
    formatting.
    Returns list of split runs starting with original {run}.
    """
    txt = run.text
    txt_len = len(txt)
    if not all(isinstance(i, int) for i in split_by):
        raise ValueError("Split positions must be integer numbers")
    split_list = [i if i >= 0 else txt_len + i for i in split_by]
    if not all(split_list[j] <= split_list[j + 1]
               for j in range(len(split_list) - 1)):
        raise ValueError("Split positions must be sorted to make sense")
    if split_list[0] < 0:
        raise ValueError("A split position cannot be less than -<text length>")
    split_list.insert(0, 0)
    split_list.append(None)
    split_txts = [txt[split_list[i]:split_list[i + 1]]
                  for i in range(len(split_list) - 1)]
    run.text = split_txts[0]
    split_txts.pop(0)
    new_runs = [run]
    for next_txt in split_txts:
        new_runs.append(insert_run_after(par, new_runs[-1], next_txt))
        copy_run_format(run, new_runs[-1])

    return new_runs


diff_list = [('Deleted', '1'), ('Deleted', ' '), ('Deleted', '2'), ('Deleted', ' '), ('Deleted', '3'), ('Deleted', ' '),
             ('Deleted', '4'), ('Deleted', ' '), ('Deleted', '5')]
# 已经完成切分，现在要根据 上面的diff处理word
new_runs = []

d = Document('1.docx')

par = d.paragraphs[3]
run = par.runs[0]
new_runs = split_run_by(par, run, [3, 5])

for i in range(len(new_runs)-1):
    if i == 1:
        new_runs[i].font.strike = True
        new_runs[i].text = new_runs[i].text + "sjkghbeha"

modified_word_document_path="1212.docx"
if os.path.exists(modified_word_document_path):
    os.remove(modified_word_document_path)
d.save(modified_word_document_path)
print("已保存修改后的 Word 文档:", modified_word_document_path)
